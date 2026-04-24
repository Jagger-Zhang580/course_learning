#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
文件解析模块
支持PDF、Word、TXT等格式的简历解析
"""

import os
import re
from typing import Optional, Dict, List
from pathlib import Path

try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    print("警告: PyPDF2未安装，PDF解析将不可用")

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("警告: python-docx未安装，Word解析将不可用")

try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False


class ResumeFileParser:
    """简历文件解析器"""
    
    SUPPORTED_FORMATS = ['.txt', '.pdf', '.docx', '.doc']
    
    def __init__(self):
        self.parsers = {
            '.txt': self._parse_txt,
            '.pdf': self._parse_pdf,
            '.docx': self._parse_docx,
            '.doc': self._parse_doc,
        }
    
    def parse(self, file_path: str) -> Dict:
        """
        解析简历文件
        
        Args:
            file_path: 文件路径
            
        Returns:
            解析结果字典
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            return {
                "success": False,
                "error": f"文件不存在: {file_path}",
                "text": ""
            }
        
        suffix = file_path.suffix.lower()
        
        if suffix not in self.SUPPORTED_FORMATS:
            return {
                "success": False,
                "error": f"不支持的文件格式: {suffix}",
                "text": ""
            }
        
        parser = self.parsers.get(suffix)
        if parser is None:
            return {
                "success": False,
                "error": f"没有可用的解析器: {suffix}",
                "text": ""
            }
        
        try:
            text = parser(str(file_path))
            return {
                "success": True,
                "text": text,
                "file_name": file_path.name,
                "file_format": suffix,
                "file_size": file_path.stat().st_size
            }
        except Exception as e:
            return {
                "success": False,
                "error": f"解析失败: {str(e)}",
                "text": ""
            }
    
    def _parse_txt(self, file_path: str) -> str:
        """解析TXT文件"""
        encodings = ['utf-8', 'gbk', 'gb2312', 'latin-1']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue
        
        raise ValueError("无法解码文件")
    
    def _parse_pdf(self, file_path: str) -> str:
        """解析PDF文件"""
        text = ""
        
        # 优先使用pdfplumber（更好的中文支持）
        if PDFPLUMBER_AVAILABLE:
            try:
                import pdfplumber
                with pdfplumber.open(file_path) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
                if text.strip():
                    return self._clean_text(text)
            except Exception:
                pass
        
        # 备用方案：使用PyPDF2
        if PDF_AVAILABLE:
            try:
                with open(file_path, 'rb') as f:
                    reader = PyPDF2.PdfReader(f)
                    for page in reader.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
                return self._clean_text(text)
            except Exception as e:
                raise ValueError(f"PDF解析失败: {e}")
        
        raise ValueError("没有可用的PDF解析库")
    
    def _parse_docx(self, file_path: str) -> str:
        """解析DOCX文件"""
        if not DOCX_AVAILABLE:
            raise ValueError("python-docx库未安装")
        
        doc = Document(file_path)
        text = ""
        
        # 提取段落
        for para in doc.paragraphs:
            if para.text.strip():
                text += para.text + "\n"
        
        # 提取表格内容
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    text += " | ".join(row_text) + "\n"
        
        return self._clean_text(text)
    
    def _parse_doc(self, file_path: str) -> str:
        """解析DOC文件（旧版Word格式）"""
        # 尝试使用textract（如果可用）
        try:
            import textract
            text = textract.process(file_path).decode('utf-8')
            return self._clean_text(text)
        except ImportError:
            pass
        
        # 提示用户转换格式
        raise ValueError(
            "旧版DOC格式需要额外库支持。"
            "建议将文件转换为DOCX格式，或安装textract库。"
        )
    
    def _clean_text(self, text: str) -> str:
        """清理文本"""
        # 移除多余空白
        text = re.sub(r'\n\s*\n', '\n\n', text)
        text = re.sub(r' +', ' ', text)
        
        # 移除特殊字符但保留中文标点
        text = re.sub(r'[^\w\s\u4e00-\u9fff.,;:!?，。；：！？、\n-]', '', text)
        
        return text.strip()
    
    def parse_from_bytes(self, content: bytes, file_format: str) -> Dict:
        """
        从字节流解析文件（用于上传文件）
        
        Args:
            content: 文件内容字节流
            file_format: 文件格式（如'.pdf'）
            
        Returns:
            解析结果
        """
        import tempfile
        
        if file_format.lower() not in self.SUPPORTED_FORMATS:
            return {
                "success": False,
                "error": f"不支持的文件格式: {file_format}",
                "text": ""
            }
        
        # 创建临时文件
        with tempfile.NamedTemporaryFile(
            suffix=file_format,
            delete=False
        ) as tmp:
            tmp.write(content)
            tmp_path = tmp.name
        
        try:
            result = self.parse(tmp_path)
            return result
        finally:
            # 清理临时文件
            os.unlink(tmp_path)


class BatchResumeParser:
    """批量简历解析器"""
    
    def __init__(self):
        self.parser = ResumeFileParser()
    
    def parse_directory(self, directory: str) -> List[Dict]:
        """
        解析目录下的所有简历文件
        
        Args:
            directory: 目录路径
            
        Returns:
            解析结果列表
        """
        results = []
        directory = Path(directory)
        
        if not directory.is_dir():
            return results
        
        for file_path in directory.iterdir():
            if file_path.suffix.lower() in ResumeFileParser.SUPPORTED_FORMATS:
                result = self.parser.parse(str(file_path))
                result['file_path'] = str(file_path)
                results.append(result)
        
        return results
    
    def parse_multiple(self, file_paths: List[str]) -> List[Dict]:
        """
        解析多个文件
        
        Args:
            file_paths: 文件路径列表
            
        Returns:
            解析结果列表
        """
        results = []
        
        for file_path in file_paths:
            result = self.parser.parse(file_path)
            result['file_path'] = file_path
            results.append(result)
        
        return results


def extract_resume_sections(text: str) -> Dict[str, str]:
    """
    提取简历的各个部分
    
    Args:
        text: 简历文本
        
    Returns:
        各部分的内容
    """
    sections = {
        "基本信息": "",
        "教育背景": "",
        "工作经验": "",
        "项目经验": "",
        "专业技能": "",
        "自我评价": "",
        "其他": ""
    }
    
    # 定义各部分的关键词模式
    patterns = {
        "基本信息": r"(个人信息|基本信息|个人简介|姓名|联系方式)",
        "教育背景": r"(教育背景|教育经历|学历|毕业院校|学校)",
        "工作经验": r"(工作经验|工作经历|职业经历|工作)",
        "项目经验": r"(项目经验|项目经历|项目)",
        "专业技能": r"(专业技能|技能|技术栈|技术能力)",
        "自我评价": r"(自我评价|个人评价|自我总结|个人总结)",
    }
    
    # 按行分割
    lines = text.split('\n')
    current_section = "其他"
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        # 检查是否是新的部分标题
        section_found = False
        for section, pattern in patterns.items():
            if re.search(pattern, line, re.IGNORECASE):
                current_section = section
                section_found = True
                break
        
        # 添加内容到当前部分
        if not section_found:
            sections[current_section] += line + "\n"
    
    # 清理各部分
    for key in sections:
        sections[key] = sections[key].strip()
    
    return sections


# 示例用法
if __name__ == "__main__":
    parser = ResumeFileParser()
    
    # 测试TXT解析
    test_txt = "test_resume.txt"
    with open(test_txt, 'w', encoding='utf-8') as f:
        f.write("""
个人简历
姓名：张三
电话：13800138000
邮箱：zhangsan@example.com

教育背景
2015-2019  北京大学  计算机科学与技术  本科

工作经验
2019-至今  某科技公司  Python开发工程师
- 负责公司核心业务系统的开发和维护
- 使用Django框架开发Web应用
- 优化数据库查询性能

专业技能
- 熟练掌握Python编程语言
- 熟悉Django、Flask框架
- 了解MySQL、Redis数据库
- 熟悉Git版本控制
        """)
    
    result = parser.parse(test_txt)
    print(f"解析结果: {result['success']}")
    print(f"文本长度: {len(result['text'])} 字符")
    
    # 提取各部分
    sections = extract_resume_sections(result['text'])
    print("\n简历各部分:")
    for section, content in sections.items():
        if content:
            print(f"\n【{section}】")
            print(content[:100] + "..." if len(content) > 100 else content)
    
    # 清理测试文件
    os.unlink(test_txt)